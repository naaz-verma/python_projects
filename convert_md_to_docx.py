"""
Generate course structure Word documents matching the WorldWithWeb template.
Reads structured data directly -- no markdown parsing needed.
Matches exact styling from 01_Python_Starter.docx.

Run: python convert_md_to_docx.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xCC, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

LOGO_PATH = 'logo.jpg'


def create_course_doc(data, output_path):
    """Create a course structure .docx from structured data."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Logo ---
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(LOGO_PATH, height=Inches(0.72))

    # --- Title ---
    p = doc.add_paragraph()
    run = p.add_run(data['title'])
    run.font.size = Pt(26)
    run.font.color.rgb = RED
    run.bold = True

    # --- Track ---
    p = doc.add_paragraph()
    run = p.add_run(f"Track: {data['track']}")
    run.font.color.rgb = GRAY

    # --- Course Overview ---
    h = doc.add_heading('Course Overview', level=1)
    p = doc.add_paragraph(data['overview'])

    # --- Course Details ---
    doc.add_heading('Course Details', level=1)
    for detail in data['details']:
        p = doc.add_paragraph()
        run = p.add_run(detail)
        run.bold = True
        run.font.color.rgb = DARK_GRAY

    # --- What Students Will Learn ---
    doc.add_heading('What Students Will Learn', level=1)
    for item in data['learn']:
        doc.add_paragraph(item, style='List Bullet')

    # --- Course Syllabus ---
    doc.add_heading('Course Syllabus', level=1)
    for section_name, bullets in data['syllabus']:
        h2 = doc.add_heading(section_name, level=2)
        for run in h2.runs:
            run.font.color.rgb = RED
        for bullet in bullets:
            doc.add_paragraph(bullet, style='List Bullet')

    # --- Projects Students Build ---
    doc.add_heading('Projects Students Build', level=1)
    for proj_name, proj_desc in data['projects']:
        p = doc.add_paragraph()
        run = p.add_run(proj_name)
        run.bold = True
        run.font.color.rgb = DARK_GRAY
        doc.add_paragraph(proj_desc)

    # --- Portfolio Outcome ---
    doc.add_heading('Portfolio Outcome', level=1)
    doc.add_paragraph('Students graduate with:')
    for item in data['outcome']:
        doc.add_paragraph(item, style='List Bullet')

    # --- Footer ---
    p = doc.add_paragraph()
    run1 = p.add_run('WorldWithWeb')
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.color.rgb = RED
    run2 = p.add_run('  |  worldwithweb.com/courses  |  +91 9872606864')
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    doc.save(output_path)
    print(f'  Created: {output_path}')


# ============================================================
# COURSE DATA
# ============================================================

courses = [
    {
        'filename': '14_Web_Dev_MERN',
        'title': 'Web Development - MERN Stack',
        'track': 'Development',
        'overview': 'Master modern full-stack web development using the MERN stack (MongoDB, Express, React, Node.js). Students build frontend and backend applications from scratch, learn TypeScript and modern tooling, and deploy 4 production-ready web apps \u2013 all showcased on GitHub and LinkedIn.',
        'details': [
            'Duration: 6 Months',
            'Fee: Rs. 60,000',
            'Level: Beginner to Advanced',
            'Mode: Practical + Project Based',
            'Tools: HTML, CSS, JavaScript, React, Node.js, Express, MongoDB, TypeScript, Tailwind CSS, Git, GitHub, AI Tools',
        ],
        'learn': [
            'HTML5 \u2013 semantic elements, forms, accessibility',
            'CSS3 \u2013 flexbox, grid, responsive design, Tailwind CSS',
            'JavaScript ES6+ \u2013 async/await, closures, array methods, modules',
            'React \u2013 hooks, routing, state management, component patterns',
            'Node.js & Express \u2013 REST APIs, middleware, authentication',
            'MongoDB & Mongoose \u2013 data modeling, CRUD, aggregation',
            'TypeScript \u2013 types, interfaces, generics with React & Express',
            'JWT authentication, input validation, security best practices',
            'Deployment \u2013 Vercel, Railway, MongoDB Atlas',
            'Git, GitHub, and professional development workflows',
        ],
        'syllabus': [
            ('Month 1: Frontend Foundations', [
                'HTML5 \u2013 document structure, forms, tables, semantic elements',
                'CSS3 \u2013 flexbox, grid, responsive design, media queries',
                'JavaScript fundamentals \u2013 variables, functions, loops, DOM',
                'JavaScript advanced \u2013 promises, async/await, fetch API, modules',
                'Tailwind CSS for styling',
                'Project: Personal Portfolio Website',
            ]),
            ('Month 2: React & Modern Frontend', [
                'React with Vite \u2013 components, JSX, props, state',
                'React Hooks \u2013 useState, useEffect, useContext, useRef',
                'React Router \u2013 nested routes, dynamic routes, protected routes',
                'Forms with React Hook Form and Zod validation',
                'State management \u2013 Zustand and TanStack Query',
                'Project: Task Management App (React)',
            ]),
            ('Month 3-4: Backend & Database', [
                'Node.js \u2013 modules, npm, environment variables',
                'Express.js \u2013 routing, middleware, error handling',
                'REST API design \u2013 CRUD, HTTP methods, status codes',
                'JWT authentication \u2013 signup, login, token verification',
                'MongoDB & Mongoose \u2013 schemas, models, validation, Atlas',
                'File uploads with Multer and Cloudinary',
                'Project: Full-Stack E-Commerce Platform',
            ]),
            ('Month 5-6: Full-Stack, TypeScript & Deployment', [
                'Connecting React frontend to Express backend',
                'TypeScript with React and Express',
                'Next.js basics \u2013 SSR, SSG, file-based routing',
                'Testing basics \u2013 Vitest, React Testing Library',
                'Deployment \u2013 Vercel, Railway, GitHub Actions CI/CD',
                'Project: Capstone App (student\u2019s choice, Next.js)',
                'GitHub portfolio finalization and LinkedIn showcase',
            ]),
        ],
        'projects': [
            ('Project 1: Personal Portfolio Website', 'A responsive portfolio site built with HTML, CSS, and Tailwind \u2013 showcasing skills, projects, and contact information. Deployed on GitHub Pages.'),
            ('Project 2: Task Management App', 'A React app with full CRUD functionality, filters, categories, and persistent state. Built with Vite, React Router, and Zustand.'),
            ('Project 3: Full-Stack E-Commerce Platform', 'Product catalog, shopping cart, checkout, user authentication, and admin dashboard. React frontend + Express API + MongoDB database.'),
            ('Project 4: Capstone Application', 'A full-stack application of the student\u2019s choice built with Next.js and TypeScript. Deployed on Vercel with CI/CD pipeline.'),
        ],
        'outcome': [
            '4 deployed web projects with live URLs',
            'GitHub repositories with documentation',
            'LinkedIn project showcase',
            'Full-stack development and deployment experience',
            'Resume-ready web developer profile',
        ],
    },
    {
        'filename': '15_Digital_Marketing',
        'title': 'Digital Marketing',
        'track': 'Digital Skills',
        'overview': 'Learn digital marketing from strategy to execution. Students master SEO, Google Ads, Meta Ads, social media marketing, content creation, email marketing, and analytics \u2013 running real campaigns with real budgets and earning industry certifications along the way.',
        'details': [
            'Duration: 3 Months',
            'Fee: Rs. 30,000',
            'Level: Beginner',
            'Mode: Practical + Project Based',
            'Tools: Google Ads, Meta Ads, GA4, Search Console, Canva, WordPress, ChatGPT, Zapier, Kit/Brevo, Looker Studio',
        ],
        'learn': [
            'Digital marketing strategy and marketing funnels',
            'WordPress website setup and management',
            'SEO \u2013 on-page, off-page, technical, local SEO',
            'Google Ads \u2013 search campaigns, Performance Max, bidding strategies',
            'Meta Ads \u2013 Advantage+ campaigns, creative strategy, A/B testing',
            'Content marketing and copywriting (PAS, AIDA frameworks)',
            'Social media marketing \u2013 Instagram, LinkedIn, YouTube',
            'Email marketing \u2013 list building, automation, A/B testing',
            'Google Analytics 4 \u2013 event tracking, reports, conversions',
            'Marketing automation with Zapier',
            'AI tools for content creation and campaign optimization',
        ],
        'syllabus': [
            ('Month 1: Foundations \u2013 Website, SEO & Strategy', [
                'Digital marketing landscape and marketing funnels',
                'Customer journey, STP, and channel overview',
                'WordPress setup \u2013 themes, plugins, pages, posts',
                'SEO fundamentals \u2013 keyword research, on-page, technical SEO',
                'Google Search Console setup and analysis',
                'Off-page SEO, local SEO, and Google Business Profile',
                'Project: WordPress Blog/Business Site + SEO Audit Report',
            ]),
            ('Month 2: Content, Social Media & Paid Ads', [
                'Content strategy \u2013 pillars, clusters, calendar planning',
                'Copywriting frameworks \u2013 PAS, AIDA',
                'Social media marketing \u2013 Instagram, LinkedIn, YouTube strategy',
                'Content creation with Canva and AI tools',
                'Google Ads \u2013 search campaigns, Performance Max, conversion tracking',
                'Meta Ads \u2013 Advantage+ campaigns, audience targeting, creative strategy',
                'Project: 30-Day Social Media Campaign + Real Ad Campaigns',
            ]),
            ('Month 3: Email, Analytics & Capstone', [
                'Email marketing \u2013 list building, automation, A/B testing (Kit/Brevo)',
                'Google Analytics 4 \u2013 event tracking, reports, explorations',
                'Google Tag Manager and UTM tracking',
                'Looker Studio dashboards and reporting',
                'Influencer marketing and affiliate marketing fundamentals',
                'Marketing automation with Zapier',
                'Certification prep \u2013 Google Ads, GA4, HubSpot Inbound',
                'Capstone: Complete Digital Marketing Case Study for a real business',
            ]),
        ],
        'projects': [
            ('Project 1: WordPress Blog/Business Site', 'A live WordPress website with 5 pages, SEO-optimized content, essential plugins, and Google Analytics integration.'),
            ('Project 2: SEO Audit Report', 'A professional audit of a real website with findings, recommendations, and actionable fixes \u2013 delivered as a PDF report.'),
            ('Project 3: Social Media Campaign + Ad Campaigns', 'A 30-day content calendar with 15 designed posts, plus real Google Ads and Meta Ads campaigns run with Rs. 500-1,000 budgets.'),
            ('Project 4: Complete Marketing Case Study', 'End-to-end digital marketing strategy for a real business \u2013 SEO, social media, paid ads, email, and analytics with documented results.'),
        ],
        'outcome': [
            '4 marketing projects with real data and results',
            '1 live WordPress website',
            'Google Ads Search + GA4 + HubSpot Inbound certifications',
            'LinkedIn portfolio with campaign showcases',
            'Resume-ready digital marketing profile',
        ],
    },
    {
        'filename': '16_MS_Office_Basics',
        'title': 'MS Office Basics',
        'track': 'Digital Skills',
        'overview': 'Learn Microsoft Office through real-world projects, not feature tutorials. Students build professional documents, data dashboards, and pitch decks \u2013 every session produces a deliverable they keep. Includes Google Workspace cross-platform literacy.',
        'details': [
            'Duration: 1.5 Months',
            'Fee: Rs. 10,500',
            'Level: Beginner',
            'Mode: Practical + Project Based',
            'Tools: Microsoft Word, Excel, PowerPoint, Outlook, Google Workspace, AI Tools',
        ],
        'learn': [
            'Microsoft Word \u2013 professional formatting, styles, TOC, mail merge',
            'Microsoft Excel \u2013 formulas, XLOOKUP, pivot tables, charts',
            'Microsoft PowerPoint \u2013 slide design, Morph transitions, Slide Master',
            'Microsoft Outlook \u2013 email etiquette, inbox management, calendar',
            'Google Workspace \u2013 Docs, Sheets, Slides (cross-platform literacy)',
            'Document collaboration and Track Changes',
            'Data analysis and visualization basics',
            'Professional resume and cover letter creation',
        ],
        'syllabus': [
            ('Week 1-2: Microsoft Word', [
                'Interface, navigation, keyboard shortcuts',
                'Formatting \u2013 styles, fonts, spacing, margins, page layout',
                'Headers, footers, page numbers, lists, tables',
                'Table of Contents, headings, citations, bibliography',
                'Track Changes, comments, collaboration',
                'Mail Merge \u2013 letters, labels, emails',
                'Project: Professional Resume + Cover Letter + Business Report',
            ]),
            ('Week 3-4: Microsoft Excel', [
                'Data entry, formatting, number formats',
                'Formulas \u2013 SUM, AVERAGE, COUNTIF, SUMIF, IF, nested IF',
                'XLOOKUP and VLOOKUP',
                'Conditional formatting, data validation, sorting, filtering',
                'Pivot Tables \u2013 creating, formatting, grouping',
                'Charts \u2013 column, bar, line, pie, combo charts',
                'Project: Personal Budget Tracker + Sales Dashboard',
            ]),
            ('Week 5-6: PowerPoint, Outlook & Capstone', [
                'Slide design principles \u2013 less text, more visuals',
                'Templates, themes, Slide Master, SmartArt',
                'Morph transitions and tasteful animations',
                'Outlook \u2013 email etiquette, folders, rules, calendar',
                'Google Workspace overview \u2013 Docs, Sheets, Slides',
                'AI tools awareness \u2013 Copilot, ChatGPT for drafting',
                'Capstone: Mini Business Plan (Word + Excel + PowerPoint)',
            ]),
        ],
        'projects': [
            ('Project 1: Professional Resume & Cover Letter', 'A formatted resume in two versions (student and experienced template) with matching cover letter. Uploaded to LinkedIn.'),
            ('Project 2: Business Report', 'A 5-8 page formatted report with auto-generated Table of Contents, headers, images, and citations.'),
            ('Project 3: Budget Tracker & Sales Dashboard', 'An Excel workbook with formulas, conditional formatting, pivot tables, and charts for data analysis.'),
            ('Project 4: Mini Business Plan (Capstone)', 'An integrated project: Word document (business proposal), Excel workbook (financial projections), and PowerPoint deck (investor pitch).'),
        ],
        'outcome': [
            '4 real projects across Word, Excel, and PowerPoint',
            'Professional resume ready to use immediately',
            'Google Drive portfolio folder with all work',
            'Cross-platform literacy (Microsoft + Google)',
            'Resume-ready with MS Office proficiency',
        ],
    },
    {
        'filename': '17_Advanced_Excel',
        'title': 'Advanced Excel',
        'track': 'Digital Skills',
        'overview': 'Go beyond formulas and pivot tables. Students master Power Query for data transformation, Power Pivot for data modeling, advanced dashboards, and VBA/Macros for automation \u2013 graduating with an analytics portfolio that demonstrates job-ready Excel expertise.',
        'details': [
            'Duration: 1 Month',
            'Fee: Rs. 12,000',
            'Level: Intermediate to Advanced',
            'Mode: Practical + Project Based',
            'Tools: Microsoft Excel 365, Power Query, Power Pivot, VBA Editor, AI Tools',
        ],
        'learn': [
            'Advanced formulas \u2013 XLOOKUP, INDEX-MATCH, SUMIFS, dynamic arrays',
            'Dynamic array functions \u2013 FILTER, SORT, SORTBY, UNIQUE',
            'Text, date, and logical functions for real-world data',
            'Power Query \u2013 ETL, data cleaning, merging, appending',
            'Power Pivot \u2013 data modeling, relationships, star schema',
            'DAX basics \u2013 CALCULATE, SUMX, AVERAGEX, DISTINCTCOUNT',
            'Interactive dashboard design with slicers and KPI cards',
            'VBA/Macros \u2013 recording, loops, conditions, UserForms',
            'Report automation \u2013 one-click report generation',
        ],
        'syllabus': [
            ('Week 1: Advanced Formulas & Dynamic Arrays', [
                'XLOOKUP, INDEX-MATCH, SUMIFS, COUNTIFS, nested IF',
                'Text functions \u2013 LEFT, RIGHT, MID, TRIM, TEXTJOIN',
                'Date functions \u2013 DATEDIF, EOMONTH, WORKDAY, NETWORKDAYS',
                'IFS, SWITCH, IFERROR, IFNA',
                'Dynamic arrays \u2013 FILTER, SORT, SORTBY, UNIQUE, SEQUENCE',
                'Named Ranges and dependent dropdown validation',
                'Project: HR Analytics Workbook',
            ]),
            ('Week 2: Power Query & Data Transformation', [
                'Power Query \u2013 importing CSV, Excel, web, folder data',
                'Data cleaning \u2013 remove duplicates, split columns, change types',
                'Transformations \u2013 unpivot, pivot, merge queries, append',
                'Custom columns and conditional columns',
                'Refreshing queries and scheduled updates',
                'Project: Multi-Source Data Consolidation',
            ]),
            ('Week 3: Power Pivot, DAX & Dashboards', [
                'Power Pivot \u2013 data model, relationships between tables',
                'DAX \u2013 CALCULATE, SUMX, AVERAGEX, DISTINCTCOUNT, RELATED',
                'Measures vs. calculated columns',
                'Advanced Pivot Tables with slicers and timelines',
                'Dashboard design \u2013 charts, KPI cards, sparklines',
                'Project: Interactive Business Dashboard',
            ]),
            ('Week 4: VBA/Macros & Capstone', [
                'Macro recording \u2013 record, edit, assign to button',
                'VBA basics \u2013 variables, loops, If/Then/Else',
                'Working with Range, Cells, Worksheets objects',
                'UserForms for data entry',
                'Automating reports \u2013 one-click generation and PDF export',
                'Project: Automated Report Generator (Capstone)',
            ]),
        ],
        'projects': [
            ('Project 1: HR Analytics Workbook', 'Employee data analysis with advanced formulas \u2013 age calculations, department lookups, salary categorization, and attendance tracking.'),
            ('Project 2: Multi-Source Data Consolidation', 'Import, clean, and combine 3+ data files from different sources into a single analysis-ready dataset using Power Query.'),
            ('Project 3: Interactive Business Dashboard', 'A full dashboard with slicers, KPI cards, multiple chart types, and a Power Pivot data model from a real business dataset.'),
            ('Project 4: Automated Report Generator', 'A VBA macro that takes raw data, processes it, creates a formatted report with charts, and exports to PDF \u2013 all with one button click.'),
        ],
        'outcome': [
            '4 advanced Excel projects (interview-ready)',
            'Power Query + Power Pivot + DAX proficiency',
            'VBA automation capability',
            'Google Drive portfolio with dashboard screenshots',
            'Resume-ready with advanced Excel skills',
        ],
    },
    {
        'filename': '18_AI_Tools_Mastery',
        'title': 'AI Tools Mastery',
        'track': 'Digital Skills',
        'overview': 'Learn to use AI as a professional tool across writing, research, design, video, and automation. No coding required. Students master prompt engineering, build real AI-powered workflows, and create publishable content using the latest AI tools.',
        'details': [
            'Duration: 1.5 Months',
            'Fee: Rs. 12,000',
            'Level: Beginner',
            'Mode: Practical + Project Based',
            'Tools: ChatGPT, Claude, Gemini, Perplexity, DALL-E, Midjourney, CapCut, Canva AI, Zapier, Bolt.new',
        ],
        'learn': [
            'AI fundamentals \u2013 how LLMs work, capabilities, limitations',
            'Prompt engineering \u2013 role setting, chain-of-thought, few-shot',
            'AI for research \u2013 Perplexity, NotebookLM, fact-checking',
            'AI for writing \u2013 ChatGPT, Claude, Gemini for content creation',
            'AI image generation \u2013 DALL-E, Ideogram, Midjourney, Canva AI',
            'AI video creation \u2013 CapCut, Runway, HeyGen',
            'AI for presentations \u2013 Gamma, PowerPoint Copilot',
            'No-code automation \u2013 Zapier and Make workflows',
            'No-code app building \u2013 Bolt.new, v0.dev',
            'AI ethics, bias awareness, and responsible use',
        ],
        'syllabus': [
            ('Week 1-2: AI Fundamentals, Prompting & Research', [
                'What AI is \u2013 tokens, probabilities, context windows, hallucinations',
                'AI ethics \u2013 plagiarism, bias, privacy, responsible use',
                'Prompt engineering \u2013 role setting, context, constraints, refinement',
                'Comparing outputs across ChatGPT, Claude, and Gemini',
                'Deep research with Perplexity and NotebookLM',
                'AI-assisted writing and editing workflows',
                'Project: Prompt Library + AI Research Report',
            ]),
            ('Week 3-4: AI for Visual & Video Content', [
                'Image generation \u2013 DALL-E, Ideogram, Midjourney, Leonardo AI',
                'Prompt engineering for images \u2013 style, composition, aspect ratio',
                'Canva AI \u2013 Magic Design, background removal, text-to-image',
                'Video editing with CapCut \u2013 cuts, transitions, auto-captions',
                'AI video tools \u2013 Runway, HeyGen avatars',
                'Audio tools \u2013 ElevenLabs text-to-speech',
                'Project: Visual Content Portfolio + Video Content Package',
            ]),
            ('Week 5-6: Automation, Apps & Capstone', [
                'Presentations with Gamma and Beautiful.ai',
                'AI in Microsoft Office and Google Workspace',
                'Automation with Zapier/Make \u2013 form to email to sheet workflows',
                'No-code app building with Bolt.new and v0.dev',
                'AI for career \u2013 resume optimization, interview prep',
                'How to evaluate and stay current with new AI tools',
                'Capstone: AI-Powered Business Solution (case study)',
            ]),
        ],
        'projects': [
            ('Project 1: Prompt Library', 'A curated collection of 30+ tested prompts across categories (writing, research, analysis, brainstorming) with before/after examples.'),
            ('Project 2: AI Research Report', 'A 10-page research report using multiple AI tools, with documented process showing which tools were used at each step.'),
            ('Project 3: Visual & Video Content Portfolio', '10+ AI-generated images for real use cases plus 3 short-form videos with AI-generated elements.'),
            ('Project 4: AI Business Solution (Capstone)', 'An end-to-end solution to a real problem using multiple AI tools \u2013 documented as a professional case study.'),
        ],
        'outcome': [
            '4 projects demonstrating AI proficiency',
            'Professional prompt library (reusable asset)',
            'Published content on LinkedIn or blog',
            'Notion/Google Drive portfolio with all work',
            'Resume-ready with AI tool proficiency',
        ],
    },
    {
        'filename': '19_Graphics_Design',
        'title': 'Graphics & Design',
        'track': 'Digital Skills',
        'overview': 'Learn graphic design from fundamentals to professional portfolio. Students master design principles, Canva, Photoshop, Illustrator, Figma, and video editing \u2013 building a Behance portfolio with 10 real projects including a complete brand identity and client simulation.',
        'details': [
            'Duration: 3 Months',
            'Fee: Rs. 30,000',
            'Level: Beginner',
            'Mode: Practical + Project Based',
            'Tools: Canva Pro, Adobe Photoshop, Adobe Illustrator, Figma, CapCut, AI Design Tools, Behance',
        ],
        'learn': [
            'Design principles \u2013 contrast, alignment, repetition, proximity',
            'Color theory, typography, and layout composition',
            'Canva \u2013 social media design, brand kits, print materials, mockups',
            'Adobe Photoshop \u2013 photo editing, retouching, compositing',
            'Adobe Illustrator \u2013 vector graphics, logos, icons, illustrations',
            'Figma \u2013 UI/UX design, prototyping, components, auto-layout',
            'Video editing with CapCut \u2013 Reels, Shorts, motion graphics',
            'AI design tools \u2013 DALL-E, Midjourney, Ideogram, Canva AI',
            'Brand identity design and guidelines documentation',
            'Client workflow \u2013 brief to delivery process',
        ],
        'syllabus': [
            ('Month 1: Design Fundamentals, Canva & Photoshop', [
                'Design principles \u2013 C.A.R.P., color theory, typography, layout',
                'Canva \u2013 social media posts, carousels, brand kit, templates',
                'Canva advanced \u2013 posters, flyers, brochures, infographics, mockups',
                'Print specifications \u2013 CMYK, resolution, bleed, trim',
                'Photoshop \u2013 layers, masks, selection tools, blending modes',
                'Photo retouching, compositing, text effects',
                'Project: Brand Identity Kit + Business Collateral + Event Poster',
            ]),
            ('Month 2: Illustrator, Figma & AI Design', [
                'Illustrator \u2013 vector vs raster, shape tools, pen tool, pathfinder',
                'Logo design \u2013 concepts, variations, brand guidelines',
                'Figma \u2013 frames, auto-layout, components, prototyping',
                'Mobile app UI design and responsive web layouts',
                'Design systems \u2013 spacing, typography scales, components',
                'AI image generation \u2013 DALL-E, Ideogram, Midjourney for design',
                'Project: Logo Design + App UI Design + AI-Enhanced Campaign',
            ]),
            ('Month 3: Video, Brand Identity & Portfolio', [
                'Video editing with CapCut \u2013 transitions, text overlays, effects',
                'Motion graphics basics and thumbnail design',
                'Complete brand identity \u2013 mood board to brand guidelines PDF',
                'Client simulation \u2013 brief, concepts, feedback, delivery',
                'Freelancing basics \u2013 pricing, proposals, contracts',
                'Behance portfolio setup with project case studies',
                'Project: Video Content Set + Brand Identity + Client Simulation',
            ]),
        ],
        'projects': [
            ('Project 1: Brand Identity Kit', 'Logo, color palette, typography, brand board, and 5 social media templates designed in Canva.'),
            ('Project 2: Business Collateral & Event Poster', 'Business card, letterhead, flyer designed in Canva, plus a composite event poster created in Photoshop.'),
            ('Project 3: Logo Design & App UI', '3 logo concepts with brand guidelines in Illustrator, plus a 5-screen mobile app prototype in Figma.'),
            ('Project 4: Complete Brand Identity (Capstone)', 'Full brand package for a real or fictional business \u2013 logo, colors, typography, guidelines PDF, social templates, mockups, and video content.'),
        ],
        'outcome': [
            '4+ design projects across print, digital, UI, and video',
            'Behance portfolio with documented case studies',
            '20+ design pieces and 3+ short-form videos',
            'Client simulation experience (brief to delivery)',
            'Resume-ready with design tool proficiency',
        ],
    },
]


def main():
    out_dir = 'Course_Structures'
    os.makedirs(out_dir, exist_ok=True)

    print('Generating course structure documents...\n')
    for course in courses:
        path = os.path.join(out_dir, f"{course['filename']}.docx")
        create_course_doc(course, path)

    print('\nDone! All documents created in Course_Structures/')


if __name__ == '__main__':
    main()
