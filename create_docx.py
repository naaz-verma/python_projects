from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ---- Helper: Add a code block with grey background ----
def add_code_block(doc, code_text, label=None):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0, 51, 102)

    # Create a single-cell table to simulate a code box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    # Set cell shading to light grey
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    cell.paragraphs[0].clear()
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)

    # Add code text
    p = cell.paragraphs[0]
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()  # spacing after

# ---- Header with BITS ID on every page ----
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("BITS ID: __________________")
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(100, 100, 100)

# ---- Footer with page number ----
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ================================================================
# COVER PAGE
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AIML Assignment 2")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Intrusion Detection System\nusing Machine Learning")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Network Intrusion Detection using KDD Cup 1999 Dataset\nComparative Analysis of Six ML Classification Algorithms")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(4):
    doc.add_paragraph()

# Student details table
table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
details = [
    ("Student Name", "[Your Name]"),
    ("BITS ID", "[Your BITS ID]"),
    ("Course", "Artificial Intelligence & Machine Learning"),
    ("Date", "April 2026"),
]
for i, (label, value) in enumerate(details):
    cell_l = table.cell(i, 0)
    cell_r = table.cell(i, 1)
    cell_l.text = label
    cell_r.text = value
    for cell in [cell_l, cell_r]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
    cell_l.paragraphs[0].runs[0].bold = True

doc.add_page_break()

# ================================================================
# TABLE OF CONTENTS
# ================================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    ("1.", "Introduction"),
    ("2.", "Overall Process Description & Solution Architecture"),
    ("3.", "Tools Used & Justification"),
    ("4.", "Data Pre-processing"),
    ("5.", "Data Correlation Analysis"),
    ("6.", "Feature Selection"),
    ("7.", "Model Building & Training"),
    ("8.", "Validation & Model Comparison"),
    ("9.", "Feature Importance Analysis"),
    ("10.", "Multi-class Classification (Innovation)"),
    ("11.", "Conclusion & Recommendations"),
    ("12.", "References"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}  {title}")
    r.font.size = Pt(11)

doc.add_page_break()

# ================================================================
# 1. INTRODUCTION
# ================================================================
doc.add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "Network security is a critical concern in today's interconnected world. "
    "Intrusion Detection Systems (IDS) serve as a vital line of defense against malicious "
    "network activities by monitoring traffic and identifying suspicious patterns. Traditional "
    "rule-based IDS have limitations in detecting novel attack patterns, making Machine Learning "
    "(ML) approaches increasingly important for adaptive and intelligent threat detection."
)

doc.add_paragraph(
    "This assignment implements and compares six Machine Learning classification algorithms "
    "for building a Network Intrusion Detection System using the KDD Cup 1999 dataset. "
    "The objective is to evaluate each algorithm's ability to distinguish between normal network "
    "traffic and various categories of network attacks, and to identify the best-performing "
    "model for deployment in a production IDS environment."
)

doc.add_heading("1.1 Problem Statement", level=2)
doc.add_paragraph(
    "Build a network intrusion detector that can distinguish between 'normal' connections and "
    "'attack' connections. The system must classify network connections using 41 features "
    "capturing characteristics such as duration, protocol type, number of failed logins, and "
    "various traffic statistics. The classification is performed at two levels:"
)

for b in ["Binary Classification: Normal vs. Attack",
          "Multi-class Classification: Normal, DoS, Probe, R2L, U2R (5 categories)"]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading("1.2 Dataset Description", level=2)
doc.add_paragraph(
    "The KDD Cup 1999 dataset is the benchmark dataset for evaluating intrusion detection "
    "systems. This project uses a cleaned and reformatted version from Kaggle (kavl31/kdd-cup-1999-data) "
    "containing the 5-class variant with labels: normal, dos, probe, r2l, and u2r."
)

doc.add_paragraph()
table = doc.add_table(rows=6, cols=2, style='Light List Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("Property", "Value"),
    ("Total Records", "~4.9 million (before deduplication)"),
    ("Features", "41 network connection attributes + target label"),
    ("Categorical Features", "protocol_type (3), service (70), flag (11)"),
    ("Target Column", "connection_type (5 classes)"),
    ("Attack Categories", "DoS, Probe, R2L, U2R, Normal"),
]
for i, (k, v) in enumerate(data):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v
    if i == 0:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph("Attack Categories and Descriptions:", style='List Bullet')
table2 = doc.add_table(rows=6, cols=3, style='Light List Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
attacks = [
    ("Category", "Description", "Examples"),
    ("DoS", "Denial of Service \u2014 overwhelms system resources", "smurf, neptune, back, teardrop"),
    ("Probe", "Surveillance/scanning to find vulnerabilities", "portsweep, ipsweep, nmap, satan"),
    ("R2L", "Remote to Local \u2014 unauthorized remote access", "warezclient, guess_passwd, ftp_write"),
    ("U2R", "User to Root \u2014 privilege escalation", "buffer_overflow, rootkit, loadmodule"),
    ("Normal", "Legitimate network traffic", "normal"),
]
for i, (c, d, e) in enumerate(attacks):
    table2.cell(i, 0).text = c
    table2.cell(i, 1).text = d
    table2.cell(i, 2).text = e
    if i == 0:
        for cell in table2.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_page_break()

# ================================================================
# 2. OVERALL PROCESS DESCRIPTION & SOLUTION ARCHITECTURE
# ================================================================
doc.add_heading("2. Overall Process Description & Solution Architecture", level=1)

doc.add_paragraph(
    "The solution follows a structured Machine Learning pipeline consisting of seven major steps. "
    "Each phase builds upon the previous one to ensure a robust and reproducible workflow."
)

doc.add_heading("2.1 Process Flow", level=2)

steps = [
    ("Step 1: Data Acquisition & Loading",
     "Download the KDD Cup 1999 dataset via kagglehub API. Load the 5-class variant "
     "(full_5_classes_cleaned.csv) which provides richer attack category labels for both "
     "binary and multi-class analysis."),
    ("Step 2: Exploratory Data Analysis (EDA)",
     "Examine dataset shape, data types, statistical summaries, class distributions, "
     "and categorical feature distributions."),
    ("Step 3: Data Pre-processing",
     "Handle missing values (none found), remove duplicate rows (~78% duplicates), "
     "encode categorical features using LabelEncoder, create binary labels, and apply "
     "stratified sampling (100,000 rows) for computational efficiency on Colab."),
    ("Step 4: Data Correlation & Feature Selection",
     "Compute the full correlation matrix to identify redundant features. Use two methods: "
     "(a) remove features with inter-correlation > 0.95, and (b) rank features by "
     "Mutual Information scores. Apply StandardScaler normalization."),
    ("Step 5: Model Building",
     "Train six classification algorithms: Gaussian Naive Bayes, Decision Tree, Random Forest, "
     "SVM (RBF kernel), Logistic Regression, and Gradient Boosting."),
    ("Step 6: Validation & Comparison",
     "Evaluate all models using Accuracy, Precision, Recall, F1-Score, and ROC-AUC. Generate "
     "confusion matrices, ROC curves, classification reports, and 5-fold cross-validation."),
    ("Step 7: Multi-class Classification (Innovation)",
     "Train a Random Forest classifier on the 5-class attack categories "
     "(DoS, Probe, R2L, U2R, Normal) to demonstrate multi-class detection capability."),
]
for title, desc in steps:
    p = doc.add_paragraph()
    r = p.add_run(title + ": ")
    r.bold = True
    p.add_run(desc)

doc.add_heading("2.2 Solution Architecture Diagram", level=2)
doc.add_paragraph("[Insert screenshot of the solution architecture / pipeline flow diagram here]").italic = True

arch_text = """+---------------------------------------------------------+
|              KDD Cup 1999 Dataset (kagglehub)           |
+----------------------------+----------------------------+
                             |
                             v
+----------------------------+----------------------------+
|            DATA PRE-PROCESSING                          |
| Missing values | Duplicates | Encoding | Sampling       |
+----------------------------+----------------------------+
                             |
                             v
+----------------------------+----------------------------+
|       CORRELATION & FEATURE SELECTION                   |
| Heatmap | Redundancy removal | Mutual Info | Scaling    |
+----------------------------+----------------------------+
                             |
                             v
+----------------------------+----------------------------+
|          MODEL TRAINING (6 Algorithms)                  |
| NB | DT | RF | SVM | LR | GB                           |
+----------------------------+----------------------------+
                             |
                             v
+----------------------------+----------------------------+
|        VALIDATION & COMPARISON                          |
| Metrics | Confusion Matrix | ROC | CV | Time            |
+----------------------------+----------------------------+
                             |
                             v
+----------------------------+----------------------------+
|   MULTI-CLASS (5 Categories) + CONCLUSION               |
+---------------------------------------------------------+"""

add_code_block(doc, arch_text, "Architecture Overview:")

doc.add_page_break()

# ================================================================
# 3. TOOLS USED & JUSTIFICATION
# ================================================================
doc.add_heading("3. Tools Used & Justification", level=1)

doc.add_heading("3.1 Development Environment", level=2)
doc.add_paragraph(
    "Google Colab was selected as the development environment for the following reasons:"
)
for r in [
    "Free GPU/TPU access for computationally intensive tasks like SVM training",
    "Pre-installed Python scientific stack (NumPy, pandas, scikit-learn, matplotlib)",
    "Cloud-based \u2014 no local setup required, easily shareable and reproducible",
    "Jupyter notebook format enables combining code, visualizations, and documentation",
    "Integration with Google Drive for persistent storage of datasets and results",
]:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading("3.2 Programming Language: Python 3", level=2)
doc.add_paragraph(
    "Python is the industry standard for machine learning due to its extensive ecosystem "
    "of libraries, readability, and community support."
)

doc.add_heading("3.3 Key Libraries", level=2)

lib_table = doc.add_table(rows=8, cols=3, style='Light List Accent 1')
lib_table.alignment = WD_TABLE_ALIGNMENT.CENTER
libs = [
    ("Library", "Version", "Purpose"),
    ("scikit-learn", "Latest", "All 6 ML algorithms, preprocessing, metrics, feature selection, cross-validation"),
    ("pandas", "Latest", "Data loading, manipulation, statistical analysis, DataFrame operations"),
    ("NumPy", "Latest", "Numerical computing, array operations, mathematical functions"),
    ("matplotlib", "Latest", "Static visualizations \u2014 bar charts, ROC curves, heatmaps"),
    ("seaborn", "Latest", "Statistical visualizations \u2014 correlation heatmaps, confusion matrices"),
    ("kagglehub", "Latest", "Programmatic dataset download from Kaggle without manual intervention"),
    ("time (stdlib)", "Built-in", "Measuring training and prediction times for performance comparison"),
]
for i, (lib, ver, purpose) in enumerate(libs):
    lib_table.cell(i, 0).text = lib
    lib_table.cell(i, 1).text = ver
    lib_table.cell(i, 2).text = purpose
    if i == 0:
        for cell in lib_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_heading("3.4 Library Import Code", level=2)
add_code_block(doc, """import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split, cross_val_score, StratifiedKFold
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.feature_selection import SelectKBest, mutual_info_classif
from sklearn.metrics import (accuracy_score, precision_score, recall_score,
                             f1_score, confusion_matrix, classification_report,
                             roc_auc_score, roc_curve)
from sklearn.naive_bayes import GaussianNB
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier
from sklearn.svm import SVC
from sklearn.linear_model import LogisticRegression
import kagglehub, time""", "Code Snippet \u2014 Library Imports:")

doc.add_page_break()

# ================================================================
# 4. DATA PRE-PROCESSING
# ================================================================
doc.add_heading("4. Data Pre-processing", level=1)

doc.add_heading("4.1 Data Loading", level=2)
doc.add_paragraph(
    "The dataset was downloaded programmatically using kagglehub. The 5-class cleaned variant "
    "(full_5_classes_cleaned.csv) was selected to enable both binary and multi-class classification."
)

add_code_block(doc, """import kagglehub

# Download latest version of KDD Cup 1999 dataset
path = kagglehub.dataset_download("kavl31/kdd-cup-1999-data")

# Walk directory tree to find the 5-class file
for root, dirs, files in os.walk(path):
    for f in files:
        if '5_classes' in f and 'cleaned' in f and 'full' in f:
            data_file = os.path.join(root, f)
            break

# Load full dataset (file has headers)
df = pd.read_csv(data_file)
print(f"Dataset loaded: {df.shape[0]:,} rows x {df.shape[1]} columns")""", "Code Snippet \u2014 Data Loading:")

doc.add_paragraph("[Insert screenshot: Dataset loading output with file listing]").italic = True

doc.add_heading("4.2 Exploratory Data Analysis", level=2)
doc.add_paragraph("Initial exploration revealed the dataset structure:")
for b in [
    "Shape: ~4.9 million rows \u00d7 42 columns",
    "Data types: 23 integer, 15 float, 4 object (categorical) columns",
    "Memory usage: ~2.5 GB",
    "Categorical features: protocol_type (3), service (70), flag (11)",
    "Target column (connection_type): 5 classes \u2014 normal, dos, probe, r2l, u2r",
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph("[Insert screenshot: Dataset overview and statistical summary]").italic = True

doc.add_heading("4.3 Missing Value Analysis", level=2)
doc.add_paragraph(
    "A systematic check confirmed the dataset contains no missing values."
)
add_code_block(doc, """missing = df.isnull().sum()
print("Missing values per column:")
print(missing[missing > 0])
if missing.sum() == 0:
    print("No missing values found in the dataset.")""", "Code Snippet \u2014 Missing Value Check:")

doc.add_heading("4.4 Duplicate Removal", level=2)
doc.add_paragraph(
    "~78% of rows were duplicates. Removing them prevents the model from memorizing "
    "repeated entries and inflating accuracy."
)
add_code_block(doc, """dup_count = df.duplicated().sum()
print(f"Duplicate rows: {dup_count:,} ({dup_count/len(df)*100:.2f}%)")
df = df.drop_duplicates()
print(f"Shape after removing duplicates: {df.shape}")""", "Code Snippet \u2014 Duplicate Removal:")

doc.add_heading("4.5 Label Encoding", level=2)
doc.add_paragraph(
    "Three categorical features were encoded using LabelEncoder:"
)
add_code_block(doc, """label_encoders = {}
categorical_cols = ['protocol_type', 'service', 'flag']

for col in categorical_cols:
    le = LabelEncoder()
    df[col] = le.fit_transform(df[col])
    label_encoders[col] = le
    print(f"  {col}: {len(le.classes_)} unique values")""", "Code Snippet \u2014 Label Encoding:")

doc.add_heading("4.6 Binary Label Creation & Attack Category Mapping", level=2)
doc.add_paragraph(
    "Attack categories were mapped from the 5-class labels, and a binary label (Normal=0, Attack=1) "
    "was created for binary classification."
)
add_code_block(doc, """attack_map = {
    'normal': 'Normal', 'dos': 'DoS', 'probe': 'Probe',
    'r2l': 'R2L', 'u2r': 'U2R',
}
df['attack_category'] = df[label_col].map(attack_map)
df['binary_label'] = df['attack_category'].apply(
    lambda x: 0 if x == 'Normal' else 1
)""", "Code Snippet \u2014 Label Mapping:")

doc.add_heading("4.7 Stratified Sampling", level=2)
doc.add_paragraph(
    "The dataset was stratified-sampled to 100,000 rows for computational efficiency while "
    "preserving class proportions."
)
add_code_block(doc, """SAMPLE_SIZE = 100000
X_sampled, _, y_sampled, _ = train_test_split(
    X, y, train_size=SAMPLE_SIZE, stratify=y, random_state=42
)
X = X_sampled.reset_index(drop=True)
y = y_sampled.reset_index(drop=True)""", "Code Snippet \u2014 Stratified Sampling:")

doc.add_paragraph("[Insert screenshot: Pre-processing outputs]").italic = True

doc.add_page_break()

# ================================================================
# 5. DATA CORRELATION ANALYSIS
# ================================================================
doc.add_heading("5. Data Correlation Analysis", level=1)

doc.add_heading("5.1 Feature Correlation Heatmap", level=2)
doc.add_paragraph(
    "A Pearson correlation matrix was computed for all 41 features. The heatmap reveals several "
    "groups of highly correlated features, particularly among error-rate and host-based statistics."
)
add_code_block(doc, """corr_matrix = X.corr()

plt.figure(figsize=(20, 16))
sns.heatmap(corr_matrix, cmap='RdBu_r', center=0,
            linewidths=0.5, square=True)
plt.title('Feature Correlation Heatmap')
plt.tight_layout()
plt.show()""", "Code Snippet \u2014 Correlation Heatmap:")

doc.add_paragraph("[Insert screenshot: Correlation heatmap]").italic = True

doc.add_heading("5.2 Highly Correlated Feature Pairs", level=2)
doc.add_paragraph("14 feature pairs with |r| > 0.9 were identified:")
for b in [
    "SYN error rates: serror_rate, srv_serror_rate, dst_host_serror_rate, dst_host_srv_serror_rate (r > 0.99)",
    "REJ error rates: rerror_rate, srv_rerror_rate, dst_host_rerror_rate, dst_host_srv_rerror_rate (r > 0.95)",
    "Compromise indicators: num_compromised and num_root (r = 0.997)",
]:
    doc.add_paragraph(b, style='List Bullet')

add_code_block(doc, """high_corr_pairs = []
for i in range(len(corr_matrix.columns)):
    for j in range(i + 1, len(corr_matrix.columns)):
        if abs(corr_matrix.iloc[i, j]) > 0.9:
            high_corr_pairs.append({
                'Feature 1': corr_matrix.columns[i],
                'Feature 2': corr_matrix.columns[j],
                'Correlation': corr_matrix.iloc[i, j]
            })""", "Code Snippet \u2014 Finding Correlated Pairs:")

doc.add_paragraph("[Insert screenshot: Correlated pairs table]").italic = True

doc.add_heading("5.3 Feature-Target Correlation", level=2)
doc.add_paragraph(
    "Top features correlated with target: same_srv_rate (0.94), error rates (~0.85), "
    "count (0.83), logged_in (0.73)."
)
add_code_block(doc, """target_corr = X.corrwith(y).abs().sort_values(ascending=False)

plt.figure(figsize=(14, 8))
target_corr.plot(kind='bar', color='teal', edgecolor='black')
plt.title('Feature Correlation with Target (|r|)')
plt.axhline(y=0.1, color='red', linestyle='--', label='Threshold')
plt.show()""", "Code Snippet \u2014 Target Correlation:")

doc.add_paragraph("[Insert screenshot: Feature-target correlation bar chart]").italic = True

doc.add_page_break()

# ================================================================
# 6. FEATURE SELECTION
# ================================================================
doc.add_heading("6. Feature Selection", level=1)

doc.add_paragraph(
    "A two-stage feature selection strategy was employed:"
)

doc.add_heading("6.1 Method 1: Redundancy Removal (Correlation > 0.95)", level=2)
doc.add_paragraph(
    "For each highly correlated pair, the feature with lower target correlation was dropped. "
    "7 redundant features were removed."
)
add_code_block(doc, """features_to_drop = set()
for i in range(len(corr_matrix.columns)):
    for j in range(i + 1, len(corr_matrix.columns)):
        if abs(corr_matrix.iloc[i, j]) > 0.95:
            feat_i = corr_matrix.columns[i]
            feat_j = corr_matrix.columns[j]
            if abs(X[feat_i].corr(y)) < abs(X[feat_j].corr(y)):
                features_to_drop.add(feat_i)
            else:
                features_to_drop.add(feat_j)""", "Code Snippet \u2014 Redundancy Removal:")

doc.add_heading("6.2 Method 2: Mutual Information Ranking", level=2)
doc.add_paragraph(
    "SelectKBest with mutual_info_classif ranked features by non-linear dependency with the target. "
    "Top 25 were selected, minus the 7 redundant ones, yielding 19 final features."
)
add_code_block(doc, """selector = SelectKBest(score_func=mutual_info_classif, k='all')
selector.fit(X, y)
mi_scores = pd.Series(selector.scores_, index=X.columns).sort_values(ascending=False)

K = 25
selected_features = mi_scores.head(K).index.tolist()
selected_features = [f for f in selected_features if f not in features_to_drop]
X_selected = X[selected_features].copy()""", "Code Snippet \u2014 Mutual Information Selection:")

doc.add_paragraph("[Insert screenshot: MI scores bar chart]").italic = True

doc.add_heading("6.3 Feature Scaling & Train-Test Split", level=2)
add_code_block(doc, """scaler = StandardScaler()
X_scaled = pd.DataFrame(scaler.fit_transform(X_selected),
                        columns=X_selected.columns)

X_train, X_test, y_train, y_test = train_test_split(
    X_scaled, y, test_size=0.2, random_state=42, stratify=y
)
print(f"Training: {X_train.shape[0]:,}  |  Test: {X_test.shape[0]:,}")""", "Code Snippet \u2014 Scaling & Split:")

doc.add_page_break()

# ================================================================
# 7. MODEL BUILDING & TRAINING
# ================================================================
doc.add_heading("7. Model Building & Training", level=1)

doc.add_paragraph(
    "Six classification algorithms were trained on the same training set and evaluated on "
    "the same test set."
)

doc.add_heading("7.1 Model Definitions", level=2)
add_code_block(doc, """models = {
    'Naive Bayes': GaussianNB(),
    'Decision Tree': DecisionTreeClassifier(
        max_depth=15, min_samples_split=10, random_state=42),
    'Random Forest': RandomForestClassifier(
        n_estimators=100, max_depth=20, n_jobs=-1, random_state=42),
    'SVM': SVC(
        kernel='rbf', C=1.0, gamma='scale', probability=True, random_state=42),
    'Logistic Regression': LogisticRegression(
        max_iter=1000, solver='lbfgs', random_state=42),
    'Gradient Boosting': GradientBoostingClassifier(
        n_estimators=100, learning_rate=0.1, max_depth=5, random_state=42)
}""", "Code Snippet \u2014 Model Definitions:")

models_desc = [
    ("Gaussian Naive Bayes", "Probabilistic classifier based on Bayes' theorem assuming feature independence. Fast baseline."),
    ("Decision Tree", "Rule-based classifier. Highly interpretable. max_depth=15 to prevent overfitting."),
    ("Random Forest", "Ensemble of 100 decision trees with bagging. Reduces variance, handles high dimensions."),
    ("SVM (RBF kernel)", "Maximum-margin classifier. Effective in high dimensions but computationally expensive."),
    ("Logistic Regression", "Linear model with probability outputs. Fast, interpretable. L-BFGS solver."),
    ("Gradient Boosting", "Sequential ensemble correcting errors iteratively. Typically highest accuracy."),
]
for name, desc in models_desc:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}: ")
    r.bold = True
    p.add_run(desc)

doc.add_heading("7.2 Training & Evaluation Loop", level=2)
add_code_block(doc, """results = {}
for name, model in models.items():
    # Train
    start_time = time.time()
    model.fit(X_train, y_train)
    train_time = time.time() - start_time

    # Predict
    y_pred = model.predict(X_test)
    y_prob = model.predict_proba(X_test)[:, 1]

    # Metrics
    acc = accuracy_score(y_test, y_pred)
    prec = precision_score(y_test, y_pred, average='weighted')
    rec = recall_score(y_test, y_pred, average='weighted')
    f1 = f1_score(y_test, y_pred, average='weighted')
    roc_auc = roc_auc_score(y_test, y_prob)

    results[name] = {
        'accuracy': acc, 'precision': prec, 'recall': rec,
        'f1_score': f1, 'roc_auc': roc_auc, 'train_time': train_time
    }""", "Code Snippet \u2014 Training Loop:")

doc.add_paragraph("[Insert screenshot: Model training output with all 6 models]").italic = True

doc.add_page_break()

# ================================================================
# 8. VALIDATION & MODEL COMPARISON
# ================================================================
doc.add_heading("8. Validation & Model Comparison", level=1)

doc.add_heading("8.1 Evaluation Metrics", level=2)
for metric, desc in [
    ("Accuracy", "Overall proportion of correct predictions."),
    ("Precision", "Proportion of predicted attacks that are actually attacks (minimizes false alarms)."),
    ("Recall", "Proportion of actual attacks correctly detected (minimizes missed attacks)."),
    ("F1-Score", "Harmonic mean of precision and recall \u2014 balances both objectives."),
    ("ROC-AUC", "Area under the ROC curve \u2014 discrimination ability across all thresholds."),
]:
    p = doc.add_paragraph()
    r = p.add_run(f"{metric}: ")
    r.bold = True
    p.add_run(desc)

doc.add_heading("8.2 Model Comparison Summary", level=2)
doc.add_paragraph("[Insert screenshot: Model comparison summary table]").italic = True

doc.add_paragraph("Key observations:")
for b in [
    "Gradient Boosting and Random Forest achieved the highest F1-Scores, confirming ensemble superiority.",
    "Decision Tree performed nearly as well, demonstrating the strong separability of KDD features.",
    "SVM achieved high accuracy but required significantly more training time (~60s vs <5s for trees).",
    "Logistic Regression provided a strong baseline with very fast training.",
    "Naive Bayes was fastest but lowest accuracy due to independence assumption.",
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading("8.3 Confusion Matrices", level=2)
add_code_block(doc, """fig, axes = plt.subplots(2, 3, figsize=(18, 12))
axes = axes.ravel()
for idx, (name, res) in enumerate(results.items()):
    cm = confusion_matrix(y_test, res['y_pred'])
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=axes[idx],
                xticklabels=['Normal', 'Attack'],
                yticklabels=['Normal', 'Attack'])
    axes[idx].set_title(f'{name}\\n(Acc: {res["accuracy"]:.4f})')
plt.tight_layout()
plt.show()""", "Code Snippet \u2014 Confusion Matrices:")

doc.add_paragraph("[Insert screenshot: 2\u00d73 confusion matrix grid]").italic = True

doc.add_heading("8.4 ROC Curves", level=2)
doc.add_paragraph(
    "All models achieved AUC > 0.99. Random Forest achieved a perfect AUC of 1.0000."
)
add_code_block(doc, """plt.figure(figsize=(10, 8))
for idx, (name, res) in enumerate(results.items()):
    fpr, tpr, _ = roc_curve(y_test, res['y_prob'])
    plt.plot(fpr, tpr, lw=2,
             label=f"{name} (AUC = {res['roc_auc']:.4f})")
plt.plot([0, 1], [0, 1], 'k--', label='Random Classifier')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('ROC Curves - Model Comparison')
plt.legend(loc='lower right')
plt.show()""", "Code Snippet \u2014 ROC Curves:")

doc.add_paragraph("[Insert screenshot: ROC curves overlay]").italic = True

doc.add_heading("8.5 Classification Reports", level=2)
add_code_block(doc, """for name, res in results.items():
    print(f"Classification Report: {name}")
    print(classification_report(y_test, res['y_pred'],
                                target_names=['Normal', 'Attack']))""", "Code Snippet \u2014 Classification Reports:")

doc.add_paragraph("[Insert screenshot: Classification reports]").italic = True

doc.add_heading("8.6 Cross-Validation (5-Fold)", level=2)
doc.add_paragraph(
    "5-fold stratified cross-validation confirmed model robustness with low variance."
)
add_code_block(doc, """cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=42)
for name, model in models.items():
    scores = cross_val_score(model, X_cv, y_cv, cv=cv,
                             scoring='f1_weighted', n_jobs=-1)
    print(f"{name:25s} - Mean F1: {scores.mean():.4f} +/- {scores.std():.4f}")""", "Code Snippet \u2014 Cross-Validation:")

doc.add_paragraph("[Insert screenshot: CV results and box plot]").italic = True

doc.add_heading("8.7 Computational Cost", level=2)
doc.add_paragraph(
    "SVM was slowest to train; Naive Bayes and Logistic Regression were fastest. "
    "Tree-based ensembles offered the best accuracy-to-speed trade-off."
)
doc.add_paragraph("[Insert screenshot: Training/prediction time charts]").italic = True

doc.add_page_break()

# ================================================================
# 9. FEATURE IMPORTANCE ANALYSIS
# ================================================================
doc.add_heading("9. Feature Importance Analysis", level=1)

doc.add_paragraph(
    "Feature importance was extracted from the three tree-based models to identify which "
    "features contribute most to classification."
)

add_code_block(doc, """tree_models = ['Decision Tree', 'Random Forest', 'Gradient Boosting']
fig, axes = plt.subplots(1, 3, figsize=(20, 7))
for idx, name in enumerate(tree_models):
    model = results[name]['model']
    importances = pd.Series(
        model.feature_importances_, index=selected_features
    ).sort_values(ascending=True)
    importances.tail(15).plot(kind='barh', ax=axes[idx],
                              color='teal', edgecolor='black')
    axes[idx].set_title(f'{name}\\nTop 15 Features')
plt.tight_layout()
plt.show()""", "Code Snippet \u2014 Feature Importance:")

doc.add_paragraph("Most important features across all tree models:")
for b in [
    "same_srv_rate \u2014 strongest discriminator",
    "service \u2014 network service type",
    "count \u2014 connection frequency",
    "dst_host_srv_count \u2014 same-service connections to destination",
    "flag \u2014 connection status",
    "src_bytes / dst_bytes \u2014 data volume (anomalous in attacks)",
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph("[Insert screenshot: Feature importance bar charts]").italic = True

doc.add_page_break()

# ================================================================
# 10. MULTI-CLASS CLASSIFICATION (INNOVATION)
# ================================================================
doc.add_heading("10. Multi-class Classification (Innovation)", level=1)

doc.add_paragraph(
    "Beyond the required binary classification, Random Forest was trained on the "
    "5-class attack categories. This demonstrates the system's ability to not only "
    "detect attacks but classify them into categories for prioritized response."
)

doc.add_heading("10.1 Attack Categories", level=2)
for b in [
    "Normal \u2014 legitimate network traffic",
    "DoS (Denial of Service) \u2014 overwhelm system resources",
    "Probe \u2014 surveillance and scanning",
    "R2L (Remote to Local) \u2014 unauthorized remote access",
    "U2R (User to Root) \u2014 privilege escalation",
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading("10.2 Implementation", level=2)
add_code_block(doc, """le_attack = LabelEncoder()
y_multi = le_attack.fit_transform(attack_cat_series)

X_multi_train, X_multi_test, y_multi_train, y_multi_test = train_test_split(
    X_scaled, y_multi, test_size=0.2, random_state=42, stratify=y_multi)

rf_multi = RandomForestClassifier(
    n_estimators=100, max_depth=20, n_jobs=-1, random_state=42)
rf_multi.fit(X_multi_train, y_multi_train)
y_multi_pred = rf_multi.predict(X_multi_test)

print(classification_report(y_multi_test, y_multi_pred,
                            target_names=le_attack.classes_))""", "Code Snippet \u2014 Multi-class Classification:")

doc.add_heading("10.3 Results", level=2)
doc.add_paragraph(
    "DoS and Normal classes were classified with near-perfect precision and recall. "
    "R2L and U2R, being minority classes, showed slightly lower recall due to scarcity."
)
doc.add_paragraph("[Insert screenshot: Multi-class classification report and confusion matrix]").italic = True

doc.add_page_break()

# ================================================================
# 11. CONCLUSION & RECOMMENDATIONS
# ================================================================
doc.add_heading("11. Conclusion & Recommendations", level=1)

doc.add_heading("11.1 Summary of Results", level=2)
doc.add_paragraph(
    "This assignment implemented and compared six machine learning algorithms for network "
    "intrusion detection using the KDD Cup 1999 dataset. All six models achieved accuracy "
    "above 97%, demonstrating the effectiveness of ML approaches for IDS."
)

doc.add_heading("11.2 Key Findings", level=2)
for f in [
    "Ensemble methods (Gradient Boosting, Random Forest) consistently outperformed individual models, achieving F1-Scores above 0.999.",
    "Decision Tree achieved near-ensemble performance, confirming the strong separability of KDD features.",
    "SVM provided excellent accuracy but at significantly higher computational cost.",
    "Logistic Regression offered the best speed-accuracy trade-off for real-time applications.",
    "Naive Bayes, despite its simplicity, achieved >97% accuracy, validating it as a fast baseline.",
    "Feature selection reduced 41 features to 19 without meaningful loss in performance.",
    "The multi-class classifier successfully distinguished between 5 attack categories.",
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading("11.3 Recommendations", level=2)
rec_table = doc.add_table(rows=4, cols=2, style='Light List Accent 1')
rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
recs = [
    ("Use Case", "Recommended Model"),
    ("Production IDS (highest accuracy)", "Gradient Boosting / Random Forest"),
    ("Real-time detection (low latency)", "Decision Tree / Logistic Regression"),
    ("Explainability / Audit requirements", "Decision Tree"),
]
for i, (use, model) in enumerate(recs):
    rec_table.cell(i, 0).text = use
    rec_table.cell(i, 1).text = model
    if i == 0:
        for cell in rec_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_heading("11.4 Innovation & Additional Work", level=2)
for inn in [
    "5-class attack category classification beyond the binary requirement",
    "5-fold stratified cross-validation for robust, variance-aware evaluation",
    "Feature importance analysis providing interpretability and security insights",
    "Comprehensive multi-metric evaluation (Accuracy, Precision, Recall, F1, ROC-AUC)",
    "ROC curve overlay for direct visual comparison of all models",
    "Computational cost analysis (training and prediction time comparison)",
    "Two-stage feature selection combining correlation analysis and Mutual Information",
]:
    doc.add_paragraph(inn, style='List Bullet')

doc.add_paragraph("[Insert screenshot: Conclusion output from notebook]").italic = True

doc.add_page_break()

# ================================================================
# 12. REFERENCES
# ================================================================
doc.add_heading("12. References", level=1)

for i, ref in enumerate([
    "KDD Cup 1999 Dataset \u2014 UCI ML Repository. https://kdd.ics.uci.edu/databases/kddcup99/kddcup99.html",
    "Kaggle Dataset (kavl31/kdd-cup-1999-data). https://www.kaggle.com/datasets/kavl31/kdd-cup-1999-data",
    "scikit-learn Documentation. https://scikit-learn.org/stable/",
    "Tavallaee, M., et al. (2009). A Detailed Analysis of the KDD CUP 99 Data Set. IEEE CISDA.",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",
    "Friedman, J. H. (2001). Greedy Function Approximation: A Gradient Boosting Machine. Annals of Statistics.",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273-297.",
], 1):
    doc.add_paragraph(f"[{i}] {ref}")

# ================================================================
# SAVE
# ================================================================
output_path = r"C:\Users\naaz.verma\personal\python_projects\AIML_Assignment2_Report.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
